"""
文档处理器
"""

import asyncio
import os
import json
import hashlib
from abc import ABC, abstractmethod
from datetime import datetime
from typing import Any, Dict, List, Optional, Union, AsyncGenerator
from pathlib import Path
from loguru import logger

from app.models.knowledge import Document, DocumentChunk


class DocumentProcessor(ABC):
    """文档处理器基类"""
    
    def __init__(self, name: str):
        self.name = name
    
    @abstractmethod
    def can_process(self, file_path: str, content_type: str) -> bool:
        """检查是否能处理指定文件"""
        pass
    
    @abstractmethod
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """处理文档"""
        pass
    
    @abstractmethod
    async def extract_text(self, file_path: str) -> str:
        """提取文本内容"""
        pass
    
    def get_file_hash(self, file_path: str) -> str:
        """计算文件哈希"""
        hash_md5 = hashlib.md5()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_md5.update(chunk)
        return hash_md5.hexdigest()


class TextProcessor(DocumentProcessor):
    """文本文档处理器"""
    
    def __init__(self):
        super().__init__("text")
        self.supported_extensions = {".txt", ".md", ".markdown", ".json", ".csv"}
    
    def can_process(self, file_path: str, content_type: str) -> bool:
        """检查是否能处理文本文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions or content_type.startswith("text/")
    
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """处理文本文档"""
        try:
            # 提取文本
            text_content = await self.extract_text(file_path)
            
            # 基本信息
            file_stats = os.stat(file_path)
            file_info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_stats.st_size,
                "file_hash": self.get_file_hash(file_path),
                "content_type": "text/plain",
                "character_count": len(text_content),
                "line_count": text_content.count('\n') + 1 if text_content else 0,
                "processing_time": datetime.utcnow().isoformat()
            }
            
            # 合并元数据
            if metadata:
                file_info.update(metadata)
            
            return {
                "content": text_content,
                "metadata": file_info
            }
            
        except Exception as e:
            logger.error(f"处理文本文档失败 {file_path}: {e}")
            raise
    
    async def extract_text(self, file_path: str) -> str:
        """提取文本内容"""
        try:
            # 尝试不同编码
            encodings = ['utf-8', 'gbk', 'gb2312', 'latin-1']
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    
                    # 处理特殊格式
                    ext = Path(file_path).suffix.lower()
                    if ext == ".json":
                        # 格式化JSON
                        try:
                            json_data = json.loads(content)
                            content = json.dumps(json_data, indent=2, ensure_ascii=False)
                        except:
                            pass  # 保持原内容
                    elif ext == ".csv":
                        # 简单处理CSV
                        lines = content.split('\n')
                        if lines:
                            content = '\n'.join(lines[:100])  # 限制行数
                    
                    return content
                    
                except UnicodeDecodeError:
                    continue
            
            # 如果所有编码都失败，返回错误
            raise ValueError(f"无法解码文件: {file_path}")
            
        except Exception as e:
            logger.error(f"提取文本失败 {file_path}: {e}")
            raise


class PDFProcessor(DocumentProcessor):
    """PDF文档处理器"""
    
    def __init__(self):
        super().__init__("pdf")
    
    def can_process(self, file_path: str, content_type: str) -> bool:
        """检查是否能处理PDF文件"""
        return Path(file_path).suffix.lower() == ".pdf" or content_type == "application/pdf"
    
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """处理PDF文档"""
        try:
            # 提取文本
            text_content = await self.extract_text(file_path)
            
            # 基本信息
            file_stats = os.stat(file_path)
            file_info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_stats.st_size,
                "file_hash": self.get_file_hash(file_path),
                "content_type": "application/pdf",
                "character_count": len(text_content),
                "processing_time": datetime.utcnow().isoformat()
            }
            
            # 尝试提取更多PDF信息
            try:
                pdf_info = await self._extract_pdf_metadata(file_path)
                file_info.update(pdf_info)
            except Exception as e:
                logger.warning(f"提取PDF元数据失败: {e}")
            
            # 合并元数据
            if metadata:
                file_info.update(metadata)
            
            return {
                "content": text_content,
                "metadata": file_info
            }
            
        except Exception as e:
            logger.error(f"处理PDF文档失败 {file_path}: {e}")
            raise
    
    async def extract_text(self, file_path: str) -> str:
        """提取PDF文本内容"""
        try:
            # 使用PyPDF2提取文本
            import PyPDF2
            
            text_content = ""
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num, page in enumerate(pdf_reader.pages):
                    try:
                        page_text = page.extract_text()
                        if page_text.strip():
                            text_content += f"\n--- 第 {page_num + 1} 页 ---\n"
                            text_content += page_text + "\n"
                    except Exception as e:
                        logger.warning(f"提取PDF第{page_num + 1}页失败: {e}")
                        continue
            
            return text_content.strip()
            
        except ImportError:
            logger.error("PyPDF2 未安装，无法处理PDF文件")
            raise ValueError("PDF处理器需要安装 PyPDF2 库")
        except Exception as e:
            logger.error(f"提取PDF文本失败 {file_path}: {e}")
            raise
    
    async def _extract_pdf_metadata(self, file_path: str) -> Dict[str, Any]:
        """提取PDF元数据"""
        try:
            import PyPDF2
            
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                metadata = {
                    "page_count": len(pdf_reader.pages)
                }
                
                # 提取PDF信息
                if pdf_reader.metadata:
                    pdf_meta = pdf_reader.metadata
                    if pdf_meta.get('/Title'):
                        metadata["title"] = str(pdf_meta['/Title'])
                    if pdf_meta.get('/Author'):
                        metadata["author"] = str(pdf_meta['/Author'])
                    if pdf_meta.get('/Subject'):
                        metadata["subject"] = str(pdf_meta['/Subject'])
                    if pdf_meta.get('/Creator'):
                        metadata["creator"] = str(pdf_meta['/Creator'])
                    if pdf_meta.get('/CreationDate'):
                        metadata["creation_date"] = str(pdf_meta['/CreationDate'])
                
                return metadata
                
        except Exception as e:
            logger.warning(f"提取PDF元数据失败: {e}")
            return {}


class WordProcessor(DocumentProcessor):
    """Word文档处理器"""
    
    def __init__(self):
        super().__init__("word")
        self.supported_extensions = {".docx", ".doc"}
    
    def can_process(self, file_path: str, content_type: str) -> bool:
        """检查是否能处理Word文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions or content_type in [
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword"
        ]
    
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """处理Word文档"""
        try:
            # 提取文本
            text_content = await self.extract_text(file_path)
            
            # 基本信息
            file_stats = os.stat(file_path)
            file_info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_stats.st_size,
                "file_hash": self.get_file_hash(file_path),
                "content_type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "character_count": len(text_content),
                "processing_time": datetime.utcnow().isoformat()
            }
            
            # 合并元数据
            if metadata:
                file_info.update(metadata)
            
            return {
                "content": text_content,
                "metadata": file_info
            }
            
        except Exception as e:
            logger.error(f"处理Word文档失败 {file_path}: {e}")
            raise
    
    async def extract_text(self, file_path: str) -> str:
        """提取Word文本内容"""
        try:
            # 使用python-docx提取文本
            from docx import Document
            
            doc = Document(file_path)
            text_content = ""
            
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content += paragraph.text + "\n"
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content += " | ".join(row_text) + "\n"
            
            return text_content.strip()
            
        except ImportError:
            logger.error("python-docx 未安装，无法处理Word文件")
            raise ValueError("Word处理器需要安装 python-docx 库")
        except Exception as e:
            logger.error(f"提取Word文本失败 {file_path}: {e}")
            raise


class HTMLProcessor(DocumentProcessor):
    """HTML文档处理器"""
    
    def __init__(self):
        super().__init__("html")
        self.supported_extensions = {".html", ".htm", ".xml"}
    
    def can_process(self, file_path: str, content_type: str) -> bool:
        """检查是否能处理HTML文件"""
        ext = Path(file_path).suffix.lower()
        return ext in self.supported_extensions or content_type.startswith("text/html")
    
    async def process(self, file_path: str, metadata: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
        """处理HTML文档"""
        try:
            # 提取文本
            text_content = await self.extract_text(file_path)
            
            # 基本信息
            file_stats = os.stat(file_path)
            file_info = {
                "file_path": file_path,
                "file_name": Path(file_path).name,
                "file_size": file_stats.st_size,
                "file_hash": self.get_file_hash(file_path),
                "content_type": "text/html",
                "character_count": len(text_content),
                "processing_time": datetime.utcnow().isoformat()
            }
            
            # 合并元数据
            if metadata:
                file_info.update(metadata)
            
            return {
                "content": text_content,
                "metadata": file_info
            }
            
        except Exception as e:
            logger.error(f"处理HTML文档失败 {file_path}: {e}")
            raise
    
    async def extract_text(self, file_path: str) -> str:
        """提取HTML文本内容"""
        try:
            # 读取HTML内容
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            
            # 使用BeautifulSoup提取文本
            try:
                from bs4 import BeautifulSoup
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # 移除脚本和样式
                for script in soup(["script", "style"]):
                    script.decompose()
                
                # 提取文本
                text_content = soup.get_text()
                
                # 清理文本
                lines = (line.strip() for line in text_content.splitlines())
                chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
                text_content = '\n'.join(chunk for chunk in chunks if chunk)
                
                return text_content
                
            except ImportError:
                # 如果没有BeautifulSoup，使用简单的正则表达式
                import re
                text_content = re.sub(r'<[^>]+>', '', html_content)
                text_content = re.sub(r'\s+', ' ', text_content).strip()
                return text_content
                
        except Exception as e:
            logger.error(f"提取HTML文本失败 {file_path}: {e}")
            raise


class DocumentProcessorRegistry:
    """文档处理器注册表"""
    
    def __init__(self):
        self.processors: List[DocumentProcessor] = []
        self._register_default_processors()
    
    def _register_default_processors(self):
        """注册默认处理器"""
        self.register_processor(TextProcessor())
        self.register_processor(PDFProcessor())
        self.register_processor(WordProcessor())
        self.register_processor(HTMLProcessor())
    
    def register_processor(self, processor: DocumentProcessor):
        """注册处理器"""
        self.processors.append(processor)
        logger.info(f"文档处理器已注册: {processor.name}")
    
    def get_processor(self, file_path: str, content_type: str = "") -> Optional[DocumentProcessor]:
        """获取处理器"""
        for processor in self.processors:
            if processor.can_process(file_path, content_type):
                return processor
        return None
    
    def list_supported_formats(self) -> Dict[str, List[str]]:
        """列出支持的格式"""
        formats = {}
        for processor in self.processors:
            if hasattr(processor, 'supported_extensions'):
                formats[processor.name] = list(processor.supported_extensions)
        return formats


# 全局处理器注册表
document_processor_registry = DocumentProcessorRegistry()