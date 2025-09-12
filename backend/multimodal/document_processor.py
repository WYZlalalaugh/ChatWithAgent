"""
文档处理器
"""

import asyncio
import logging
from typing import Dict, List, Optional, Any
import tempfile
import os
from pathlib import Path
import json
import re

try:
    import docx
    from docx.document import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    OPENPYXL_AVAILABLE = True
except ImportError:
    OPENPYXL_AVAILABLE = False

try:
    import PyPDF2
    import fitz  # PyMuPDF
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from pptx import Presentation
    PPTX_AVAILABLE = True
except ImportError:
    PPTX_AVAILABLE = False

try:
    import csv
    import pandas as pd
    PANDAS_AVAILABLE = True
except ImportError:
    PANDAS_AVAILABLE = False
    import csv

from .base import BaseMediaProcessor, ProcessingResult, MediaType


class PDFProcessor(BaseMediaProcessor):
    """PDF文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.pdf']
        self.max_file_size = 100 * 1024 * 1024  # 100MB
    
    def supports_format(self, file_extension: str) -> bool:
        return file_extension.lower() in self.supported_formats
    
    async def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """处理PDF文件"""
        if not PDF_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error="PDF processing libraries not available"
            )
        
        try:
            if not self.validate_file(file_path):
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    error="File validation failed"
                )
            
            # 处理选项
            extract_text = kwargs.get('extract_text', True)
            extract_images = kwargs.get('extract_images', False)
            extract_metadata = kwargs.get('extract_metadata', True)
            page_range = kwargs.get('page_range', None)
            
            content_parts = []
            processed_files = []
            doc_info = {}
            
            # 使用PyMuPDF处理
            doc = fitz.open(file_path)
            
            try:
                # 基本信息
                doc_info.update({
                    'page_count': doc.page_count,
                    'is_encrypted': doc.needs_pass,
                    'is_pdf': True
                })
                
                content_parts.append(f"PDF文档，共{doc.page_count}页")
                
                if doc.needs_pass:
                    content_parts.append("文档已加密")
                    return ProcessingResult(
                        success=False,
                        media_type=MediaType.DOCUMENT,
                        error="Document is encrypted",
                        metadata=doc_info
                    )
                
                # 提取元数据
                if extract_metadata:
                    metadata = await self._extract_pdf_metadata(doc)
                    doc_info.update(metadata)
                    content_parts.append(self._describe_pdf_metadata(metadata))
                
                # 确定页面范围
                if page_range:
                    start_page, end_page = page_range
                    start_page = max(0, start_page)
                    end_page = min(doc.page_count, end_page)
                else:
                    start_page, end_page = 0, doc.page_count
                
                # 提取文本
                if extract_text:
                    text_content = await self._extract_pdf_text(doc, start_page, end_page)
                    if text_content:
                        # 文本分析
                        text_analysis = self._analyze_text(text_content)
                        doc_info.update(text_analysis)
                        
                        content_parts.append(f"提取文本内容，共{text_analysis['char_count']}字符")
                        content_parts.append(f"文本预览: {text_content[:200]}...")
                
                # 提取图像
                if extract_images:
                    image_paths = await self._extract_pdf_images(doc, file_path, start_page, end_page)
                    if image_paths:
                        processed_files.extend(image_paths)
                        content_parts.append(f"提取了{len(image_paths)}张图像")
                
                return ProcessingResult(
                    success=True,
                    media_type=MediaType.DOCUMENT,
                    content='\n'.join(content_parts),
                    metadata=doc_info,
                    processed_files=processed_files
                )
                
            finally:
                doc.close()
                
        except Exception as e:
            self.logger.error(f"PDF processing error: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error=str(e)
            )
    
    async def _extract_pdf_metadata(self, doc: fitz.Document) -> Dict[str, Any]:
        """提取PDF元数据"""
        metadata = {}
        
        try:
            pdf_metadata = doc.metadata
            if pdf_metadata:
                metadata.update({
                    'title': pdf_metadata.get('title', ''),
                    'author': pdf_metadata.get('author', ''),
                    'subject': pdf_metadata.get('subject', ''),
                    'creator': pdf_metadata.get('creator', ''),
                    'producer': pdf_metadata.get('producer', ''),
                    'creation_date': pdf_metadata.get('creationDate', ''),
                    'modification_date': pdf_metadata.get('modDate', '')
                })
            
            # 文档统计
            total_chars = 0
            total_words = 0
            
            for page_num in range(doc.page_count):
                page = doc[page_num]
                text = page.get_text()
                total_chars += len(text)
                total_words += len(text.split())
            
            metadata.update({
                'total_characters': total_chars,
                'total_words': total_words,
                'average_chars_per_page': total_chars / doc.page_count if doc.page_count > 0 else 0
            })
            
        except Exception as e:
            self.logger.error(f"PDF metadata extraction error: {e}")
        
        return metadata
    
    async def _extract_pdf_text(self, doc: fitz.Document, start_page: int, end_page: int) -> str:
        """提取PDF文本"""
        text_parts = []
        
        try:
            for page_num in range(start_page, end_page):
                page = doc[page_num]
                text = page.get_text()
                if text.strip():
                    text_parts.append(text)
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"PDF text extraction error: {e}")
            return ""
    
    async def _extract_pdf_images(self, doc: fitz.Document, file_path: str, start_page: int, end_page: int) -> List[str]:
        """提取PDF图像"""
        image_paths = []
        
        try:
            for page_num in range(start_page, end_page):
                page = doc[page_num]
                image_list = page.get_images()
                
                for img_index, img in enumerate(image_list):
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    
                    if pix.n - pix.alpha < 4:  # 确保是RGB或灰度图像
                        img_path = self._get_output_path(
                            file_path, f"_page{page_num}_img{img_index}", ".png"
                        )
                        pix.save(img_path)
                        image_paths.append(img_path)
                    
                    pix = None
                    
                    if len(image_paths) >= 20:  # 限制提取的图像数量
                        break
                
                if len(image_paths) >= 20:
                    break
                    
        except Exception as e:
            self.logger.error(f"PDF image extraction error: {e}")
        
        return image_paths
    
    def _describe_pdf_metadata(self, metadata: Dict[str, Any]) -> str:
        """描述PDF元数据"""
        descriptions = []
        
        try:
            if metadata.get('title'):
                descriptions.append(f"标题: {metadata['title']}")
            
            if metadata.get('author'):
                descriptions.append(f"作者: {metadata['author']}")
            
            total_words = metadata.get('total_words', 0)
            if total_words > 0:
                descriptions.append(f"约{total_words}个词")
            
            return "文档信息: " + ", ".join(descriptions) if descriptions else "基本文档信息已提取"
            
        except Exception as e:
            self.logger.error(f"PDF metadata description error: {e}")
            return "文档信息提取完成"
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """分析文本内容"""
        analysis = {
            'char_count': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.split('\n')),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
        }
        
        # 语言检测（简单版本）
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        
        if chinese_chars > english_words * 2:
            analysis['primary_language'] = 'Chinese'
        elif english_words > chinese_chars:
            analysis['primary_language'] = 'English'
        else:
            analysis['primary_language'] = 'Mixed'
        
        return analysis
    
    def _get_output_path(self, original_path: str, suffix: str, extension: str = None) -> str:
        """获取输出文件路径"""
        path = Path(original_path)
        if extension:
            return str(path.parent / f"{path.stem}{suffix}{extension}")
        else:
            return str(path.parent / f"{path.stem}{suffix}{path.suffix}")


class WordProcessor(BaseMediaProcessor):
    """Word文档处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.docx', '.doc']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    def supports_format(self, file_extension: str) -> bool:
        return file_extension.lower() in self.supported_formats
    
    async def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """处理Word文档"""
        if not DOCX_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error="python-docx not available"
            )
        
        try:
            if not self.validate_file(file_path):
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    error="File validation failed"
                )
            
            # 处理选项
            extract_text = kwargs.get('extract_text', True)
            extract_images = kwargs.get('extract_images', False)
            analyze_structure = kwargs.get('analyze_structure', True)
            
            # 打开文档
            doc = docx.Document(file_path)
            
            content_parts = []
            processed_files = []
            doc_info = {}
            
            # 基本信息
            doc_info['paragraph_count'] = len(doc.paragraphs)
            doc_info['table_count'] = len(doc.tables)
            
            content_parts.append(f"Word文档，{doc_info['paragraph_count']}个段落，{doc_info['table_count']}个表格")
            
            # 提取文本
            if extract_text:
                text_content = await self._extract_word_text(doc)
                if text_content:
                    text_analysis = self._analyze_text(text_content)
                    doc_info.update(text_analysis)
                    
                    content_parts.append(f"提取文本内容，共{text_analysis['char_count']}字符")
                    content_parts.append(f"文本预览: {text_content[:200]}...")
            
            # 分析文档结构
            if analyze_structure:
                structure = await self._analyze_word_structure(doc)
                doc_info.update(structure)
                content_parts.append(self._describe_word_structure(structure))
            
            # 提取图像
            if extract_images:
                # Word图像提取比较复杂，这里提供基础实现
                content_parts.append("Word图像提取功能需要额外实现")
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.DOCUMENT,
                content='\n'.join(content_parts),
                metadata=doc_info,
                processed_files=processed_files
            )
            
        except Exception as e:
            self.logger.error(f"Word processing error: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error=str(e)
            )
    
    async def _extract_word_text(self, doc: DocxDocument) -> str:
        """提取Word文档文本"""
        text_parts = []
        
        try:
            # 提取段落文本
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 提取表格文本
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            self.logger.error(f"Word text extraction error: {e}")
            return ""
    
    async def _analyze_word_structure(self, doc: DocxDocument) -> Dict[str, Any]:
        """分析Word文档结构"""
        structure = {
            'headings': [],
            'styles_used': set(),
            'has_images': False,
            'has_tables': len(doc.tables) > 0,
            'hyperlinks': []
        }
        
        try:
            for paragraph in doc.paragraphs:
                style_name = paragraph.style.name
                structure['styles_used'].add(style_name)
                
                # 检测标题
                if 'Heading' in style_name or 'Title' in style_name:
                    structure['headings'].append({
                        'text': paragraph.text,
                        'style': style_name
                    })
                
                # 检测超链接
                for run in paragraph.runs:
                    if run.hyperlink:
                        structure['hyperlinks'].append(run.hyperlink.address)
            
            # 转换set为list以便JSON序列化
            structure['styles_used'] = list(structure['styles_used'])
            
        except Exception as e:
            self.logger.error(f"Word structure analysis error: {e}")
        
        return structure
    
    def _describe_word_structure(self, structure: Dict[str, Any]) -> str:
        """描述Word文档结构"""
        descriptions = []
        
        try:
            heading_count = len(structure.get('headings', []))
            if heading_count > 0:
                descriptions.append(f"{heading_count}个标题")
            
            if structure.get('has_tables'):
                descriptions.append("包含表格")
            
            if structure.get('has_images'):
                descriptions.append("包含图像")
            
            hyperlink_count = len(structure.get('hyperlinks', []))
            if hyperlink_count > 0:
                descriptions.append(f"{hyperlink_count}个超链接")
            
            return "文档结构: " + ", ".join(descriptions) if descriptions else "基本文档结构"
            
        except Exception as e:
            self.logger.error(f"Word structure description error: {e}")
            return "文档结构分析完成"
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """分析文本内容"""
        analysis = {
            'char_count': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.split('\n')),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
        }
        
        # 语言检测（简单版本）
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        
        if chinese_chars > english_words * 2:
            analysis['primary_language'] = 'Chinese'
        elif english_words > chinese_chars:
            analysis['primary_language'] = 'English'
        else:
            analysis['primary_language'] = 'Mixed'
        
        return analysis


class ExcelProcessor(BaseMediaProcessor):
    """Excel处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.xlsx', '.xls']
        self.max_file_size = 50 * 1024 * 1024  # 50MB
    
    def supports_format(self, file_extension: str) -> bool:
        return file_extension.lower() in self.supported_formats
    
    async def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """处理Excel文件"""
        if not OPENPYXL_AVAILABLE:
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error="openpyxl not available"
            )
        
        try:
            if not self.validate_file(file_path):
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    error="File validation failed"
                )
            
            # 处理选项
            max_rows = kwargs.get('max_rows', 1000)
            extract_formulas = kwargs.get('extract_formulas', False)
            
            # 打开工作簿
            wb = openpyxl.load_workbook(file_path, data_only=True)
            
            content_parts = []
            doc_info = {
                'worksheet_count': len(wb.worksheets),
                'worksheet_names': [ws.title for ws in wb.worksheets]
            }
            
            content_parts.append(f"Excel文件，{doc_info['worksheet_count']}个工作表")
            
            # 分析每个工作表
            for ws in wb.worksheets:
                ws_analysis = await self._analyze_worksheet(ws, max_rows)
                doc_info[f'worksheet_{ws.title}'] = ws_analysis
                content_parts.append(f"工作表'{ws.title}': {ws_analysis['used_rows']}行x{ws_analysis['used_cols']}列")
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.DOCUMENT,
                content='\n'.join(content_parts),
                metadata=doc_info
            )
            
        except Exception as e:
            self.logger.error(f"Excel processing error: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error=str(e)
            )
    
    async def _analyze_worksheet(self, worksheet, max_rows: int) -> Dict[str, Any]:
        """分析工作表"""
        analysis = {
            'used_rows': 0,
            'used_cols': 0,
            'cell_count': 0,
            'formula_count': 0,
            'data_types': {},
            'sample_data': []
        }
        
        try:
            # 获取使用范围
            if worksheet.max_row and worksheet.max_column:
                analysis['used_rows'] = worksheet.max_row
                analysis['used_cols'] = worksheet.max_column
            
            data_types = {}
            formula_count = 0
            sample_data = []
            
            # 采样数据分析
            sample_rows = min(max_rows, worksheet.max_row or 0)
            
            for row in worksheet.iter_rows(max_row=sample_rows, values_only=False):
                row_data = []
                for cell in row:
                    if cell.value is not None:
                        analysis['cell_count'] += 1
                        
                        # 统计数据类型
                        data_type = type(cell.value).__name__
                        data_types[data_type] = data_types.get(data_type, 0) + 1
                        
                        # 检测公式
                        if str(cell.value).startswith('='):
                            formula_count += 1
                        
                        row_data.append(str(cell.value)[:50])  # 限制长度
                
                if row_data and len(sample_data) < 5:  # 只保存前5行作为示例
                    sample_data.append(row_data)
            
            analysis['data_types'] = data_types
            analysis['formula_count'] = formula_count
            analysis['sample_data'] = sample_data
            
        except Exception as e:
            self.logger.error(f"Worksheet analysis error: {e}")
        
        return analysis


class TextProcessor(BaseMediaProcessor):
    """纯文本处理器"""
    
    def __init__(self):
        super().__init__()
        self.supported_formats = ['.txt', '.md', '.csv', '.json', '.xml', '.html', '.py', '.js', '.css']
        self.max_file_size = 10 * 1024 * 1024  # 10MB
    
    def supports_format(self, file_extension: str) -> bool:
        return file_extension.lower() in self.supported_formats
    
    async def process(self, file_path: str, **kwargs) -> ProcessingResult:
        """处理文本文件"""
        try:
            if not self.validate_file(file_path):
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    error="File validation failed"
                )
            
            # 尝试不同编码
            encodings = ['utf-8', 'utf-8-sig', 'gbk', 'gb2312', 'latin1']
            content = None
            used_encoding = None
            
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        content = f.read()
                    used_encoding = encoding
                    break
                except UnicodeDecodeError:
                    continue
            
            if content is None:
                return ProcessingResult(
                    success=False,
                    media_type=MediaType.DOCUMENT,
                    error="Unable to decode file with any supported encoding"
                )
            
            # 分析文本
            text_analysis = self._analyze_text(content)
            text_analysis['encoding'] = used_encoding
            text_analysis['file_extension'] = Path(file_path).suffix.lower()
            
            content_parts = [
                f"文本文件（{used_encoding}编码）",
                f"共{text_analysis['char_count']}字符，{text_analysis['word_count']}词，{text_analysis['line_count']}行"
            ]
            
            # 根据文件类型进行特殊处理
            file_ext = Path(file_path).suffix.lower()
            if file_ext == '.csv':
                csv_analysis = self._analyze_csv(content)
                text_analysis.update(csv_analysis)
                content_parts.append(f"CSV文件，{csv_analysis.get('column_count', 0)}列")
            
            elif file_ext == '.json':
                json_analysis = self._analyze_json(content)
                text_analysis.update(json_analysis)
                content_parts.append("JSON格式数据")
            
            elif file_ext in ['.py', '.js', '.css', '.html']:
                content_parts.append(f"代码文件（{file_ext[1:].upper()}）")
            
            # 内容预览
            preview = content[:300].replace('\n', ' ')
            content_parts.append(f"内容预览: {preview}...")
            
            return ProcessingResult(
                success=True,
                media_type=MediaType.DOCUMENT,
                content='\n'.join(content_parts),
                metadata=text_analysis
            )
            
        except Exception as e:
            self.logger.error(f"Text processing error: {e}")
            return ProcessingResult(
                success=False,
                media_type=MediaType.DOCUMENT,
                error=str(e)
            )
    
    def _analyze_text(self, text: str) -> Dict[str, Any]:
        """分析文本内容"""
        analysis = {
            'char_count': len(text),
            'word_count': len(text.split()),
            'line_count': len(text.split('\n')),
            'paragraph_count': len([p for p in text.split('\n\n') if p.strip()]),
        }
        
        # 语言检测
        chinese_chars = len(re.findall(r'[\u4e00-\u9fff]', text))
        english_words = len(re.findall(r'\b[a-zA-Z]+\b', text))
        
        if chinese_chars > english_words * 2:
            analysis['primary_language'] = 'Chinese'
        elif english_words > chinese_chars:
            analysis['primary_language'] = 'English'
        else:
            analysis['primary_language'] = 'Mixed'
        
        return analysis
    
    def _analyze_csv(self, content: str) -> Dict[str, Any]:
        """分析CSV文件"""
        try:
            lines = content.split('\n')
            if lines:
                first_line = lines[0]
                # 尝试检测分隔符
                delimiters = [',', ';', '\t', '|']
                delimiter = ','
                max_fields = 0
                
                for delim in delimiters:
                    field_count = len(first_line.split(delim))
                    if field_count > max_fields:
                        max_fields = field_count
                        delimiter = delim
                
                return {
                    'csv_delimiter': delimiter,
                    'column_count': max_fields,
                    'estimated_row_count': len([line for line in lines if line.strip()])
                }
        except:
            pass
        
        return {}
    
    def _analyze_json(self, content: str) -> Dict[str, Any]:
        """分析JSON文件"""
        try:
            data = json.loads(content)
            
            def count_items(obj, depth=0):
                if isinstance(obj, dict):
                    return {'objects': 1, 'max_depth': depth, 'keys': len(obj)}
                elif isinstance(obj, list):
                    return {'arrays': 1, 'max_depth': depth, 'items': len(obj)}
                else:
                    return {'primitives': 1, 'max_depth': depth}
            
            stats = count_items(data)
            return {
                'json_valid': True,
                'json_structure': type(data).__name__,
                **stats
            }
        except:
            return {'json_valid': False}


# 注册文档处理器
def register_document_processors(multimodal_processor):
    """注册文档处理器"""
    if PDF_AVAILABLE:
        multimodal_processor.register_processor(MediaType.DOCUMENT, PDFProcessor())
    
    if DOCX_AVAILABLE:
        multimodal_processor.register_processor(MediaType.DOCUMENT, WordProcessor())
    
    if OPENPYXL_AVAILABLE:
        multimodal_processor.register_processor(MediaType.DOCUMENT, ExcelProcessor())
    
    # 文本处理器总是可用
    multimodal_processor.register_processor(MediaType.DOCUMENT, TextProcessor())